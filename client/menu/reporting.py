import os

from docx.shared import Pt

from client.menu.extra_func import get_repair_hardware, get_good_dates_repair_hardware
from client.misc.func_with_time import time_now, calculate_time_difference, compare_dates
from client.misc.func_with_time import get_dates
import csv
from docx import Document  # Импортируем библиотеку для работы с документами Word

def docs_report(name, statistic):
    document = Document()  # Создание нового документа Word
    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(14)
    dates = get_dates(statistic)
    closed_applications = 0  # Общее количество закрытых заявок
    equipment_under_repair = 0  # Количество оборудования в ремонте
    unused_equipment = 0  # Общее время простаивания оборудования
    count_applications = 0  # Общее количество заявок
    repair_hardware = get_good_dates_repair_hardware(get_repair_hardware(), dates)


    for i in repair_hardware:
        start = i.get('start')
        end = i.get('end')
        done = i.get('done')
        if done == 1:
            closed_applications += 1

        elif done == 0:
            equipment_under_repair += 1
            if compare_dates(start, dates[0]) == 1:
                unused_equipment += calculate_time_difference(start, time_now())
            else:
                unused_equipment += calculate_time_difference(dates[0], time_now())

        count_applications += 1  # Увеличиваем общее количество заявок

    # Добавление заголовков и информации в документ
    document.add_heading(f'Отчет о работе оборудования за период с {dates[0]} по {dates[-1]}', 0)
    document.add_paragraph(f'Отчет предоставляет данные о работе оборудования за указанный выше период. ')
    document.add_paragraph(f'1.Общее количество заявок: {count_applications}')  # Общее количество заявок
    document.add_paragraph(f'2.Количество закрытых заявок: {closed_applications}')  # Количество закрытых заявок
    document.add_paragraph(f'3.Количество оборудования в ремонте: {equipment_under_repair}')  # Оборудование в ремонте
    document.add_paragraph(
        f"4.Оборудование простаивало(в часах): {unused_equipment}")  # Общее время простаивания оборудования
    a = document.add_paragraph()
    small_font_run = a.add_run(
        "Данные предоставляются автоматической системой сбора данных, которая может иметь недоработки")
    small_font_run.font.size = Pt(8)
    path = f'reports/{name}.docx'
    document.save(path)
    return 'Ваш отчет сохранился в ' + os.path.abspath(path)


def csv_report(name, statistic):
    data = get_good_dates_repair_hardware(get_repair_hardware(), get_dates(statistic))
    path = f'reports/{name}.csv'
    with open(path, 'w', newline='', encoding="utf8") as f:
        writer = csv.DictWriter(
            f, fieldnames=list(data[0].keys()),
            delimiter=';', quoting=csv.QUOTE_NONNUMERIC)
        writer.writeheader()
        for d in data:
            writer.writerow(d)

    return 'Ваш отчет сохранился в ' + os.path.abspath(path)

csv_report('drfg', 'year')